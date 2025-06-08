"""
generar_informe.py

Autor: Álvaro Cortés González  
Universidad: Escuela Politécnica Superior de Córdoba  
Trabajo Fin de Grado - Automatización en la Evaluación de Riesgos en Maquinaria a Partir de Imágenes Mediante Técnicas de IA  
Fecha: Junio 2025

Descripción:
    Este módulo se encarga de generar automáticamente un informe técnico en formato DOCX.
    El informe incluye la imagen analizada por IA, los riesgos detectados y las medidas
    correctoras propuestas según el tipo de fallo.
"""

from docx import Document
from docx.shared import Inches
import datetime
import os


def generar_informe(nombre_maquina, imagen_path, riesgos_detectados, informe_path):
    """
    Crea y guarda un documento Word (.docx) con los resultados de la evaluación.

    Parámetros:
    - nombre_maquina: nombre del equipo evaluado (ej. "Taladro").
    - imagen_path: ruta de la imagen marcada con los fallos.
    - riesgos_detectados: lista de riesgos detectados (diccionarios con clave 'clase').
    - informe_path: ruta completa donde se guardará el informe generado.
    """

    doc = Document()

    # Título del informe con la fecha
    fecha = datetime.datetime.now().strftime("%d-%m-%Y")
    titulo = f"Informe {nombre_maquina} {fecha}"
    doc.add_heading(titulo, 0)

    # Datos de identificación del equipo (sin rellenar)
    doc.add_paragraph(f"{nombre_maquina.upper()}.\n\nMarca:\n\nMod.:\n\nN.º de Serie:\n")

    # Sección de Estado Actual / Estado Definitivo con imagen
    doc.add_heading("ESTADO ACTUAL", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    row = table.rows[0]
    cell_actual, cell_definitivo = row.cells
    cell_actual.paragraphs[0].add_run("Estado Actual\n").bold = True
    cell_actual.paragraphs[0].add_run().add_picture(imagen_path, width=Inches(2.5))
    cell_definitivo.paragraphs[0].add_run("Estado Definitivo\n").bold = True
    cell_definitivo.paragraphs[0].add_run("[Sin especificar]")

    # Sección de riesgos detectados con descripciones legales
    doc.add_paragraph("\n8. RIESGOS DETECTADOS EN LA MÁQUINA.", style="Heading 1")

    descripciones_riesgo = {
        "falta_resguardo":
            "Proyección de objetos (requisito 4 del Anexo I del R.D. 1215/97) y protección contra elementos móviles (requisito 8).\n\n"
            "Existe riesgo de proyección de objetos desprendidos y de atrapamiento cuando se opera con dicho equipo.",
        "mandos_inseguros":
            "Mandos inseguros (requisito 9 del Anexo I del R.D. 1215/97).\n\n"
            "Los mandos del equipo no garantizan una activación segura y clara por parte del operador.",
        "sin_indicador_emergencia":
            "Advertencia y señalización (requisito 13 del Anexo I del R.D. 1215/97).\n\n"
            "El equipo carece de señalización visible sobre riesgos de operación o parada de emergencia.",
        "sin_senalizacion_riesgos":
            "Advertencia y señalización (requisito 13 del Anexo I del R.D. 1215/97).\n\n"
            "No están señalizados los riesgos del equipo de trabajo."
    }

    if riesgos_detectados:
        clases_ya_mostradas = set()
        for r in riesgos_detectados:
            clase = r["clase"]
            if clase in descripciones_riesgo and clase not in clases_ya_mostradas:
                doc.add_paragraph(descripciones_riesgo[clase])
                clases_ya_mostradas.add(clase)
    else:
        doc.add_paragraph("No se han detectado riesgos visibles en la máquina.")

    # Sección de medidas correctoras propuestas
    doc.add_paragraph("\n9. RELACIÓN DE MEDIDAS CORRECTORAS PROPUESTAS.", style="Heading 1")

    medidas = {
        "falta_resguardo":
            "Se deberán colocar resguardos adecuados para minimizar el riesgo de contactos con órganos en movimiento y proyección de partículas. "
            "Estos resguardos deberán contar con micro-interruptor de seguridad que detenga el equipo al ser abiertos.",
        "mandos_inseguros":
            "Se recomienda instalar mandos claramente identificables, accesibles y con parada de emergencia activa y señalizada.",
        "sin_indicador_emergencia":
            "Deberán incorporarse botones de parada de emergencia fácilmente visibles y accesibles, junto a señalización normalizada.",
        "sin_senalizacion_riesgos":
            "Se colocará señalización de advertencia visible sobre riesgos eléctricos, mecánicos y de uso de EPIs."
    }

    if riesgos_detectados:
        propuestas_ya_mostradas = set()
        for r in riesgos_detectados:
            clase = r["clase"]
            if clase in medidas and clase not in propuestas_ya_mostradas:
                doc.add_paragraph(medidas[clase])
                propuestas_ya_mostradas.add(clase)
    else:
        doc.add_paragraph("No se requiere ninguna medida correctora.")

    # Guardar el informe final
    doc.save(informe_path)
